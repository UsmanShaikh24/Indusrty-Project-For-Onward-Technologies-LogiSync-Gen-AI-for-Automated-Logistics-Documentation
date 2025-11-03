from typing import List, Dict, Any, Optional
import os
from datetime import datetime
import openai
from jinja2 import Environment, FileSystemLoader
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from ..config import settings

# Initialize Ollama client
from .ollama_client import OllamaClient
ollama = OllamaClient()
import openai
from datetime import datetime
from jinja2 import Environment, FileSystemLoader
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from ..config import settings

class DocumentGenerator:
    def __init__(self):
        self.env = Environment(loader=FileSystemLoader('templates'))
        openai.api_key = settings.OPENAI_API_KEY
        
    async def generate_route_document(self, route_data: Dict[str, Any]) -> str:
        """Generate a detailed route plan document"""
        # Create document template
        doc = Document()
        doc.add_heading(f'Route Plan: {route_data["route_name"]}', 0)
        
        # Add route summary
        doc.add_heading('Route Summary', level=1)
        doc.add_paragraph(f'Start Location: {route_data["start_location"]}')
        doc.add_paragraph(f'End Location: {route_data["end_location"]}')
        doc.add_paragraph(f'Total Distance: {route_data["total_distance"]} km')
        
        # Add turn-by-turn instructions
        doc.add_heading('Turn-by-Turn Instructions', level=1)
        for step in route_data["optimized_route"]:
            doc.add_paragraph(
                f'- {step["location"]}\n'
                f'  Arrival Time: {step["arrival_time"]}\n'
                f'  Cargo: {step["cargo_handling"]}'
            )
        
        # Add fuel stops
        doc.add_heading('Recommended Fuel Stops', level=1)
        for stop in route_data["fuel_stops"]:
            doc.add_paragraph(
                f'- Location: {stop["location"]}\n'
                f'  Distance from start: {stop["distance_from_start"]} km'
            )
        
        # Add compliance checkpoints
        doc.add_heading('Compliance Checkpoints', level=1)
        for checkpoint in route_data["compliance_checkpoints"]:
            doc.add_paragraph(
                f'- Location: {checkpoint["location"]}\n'
                f'  Type: {checkpoint["type"]}\n'
                f'  Duration: {checkpoint["duration_minutes"]} minutes'
            )
        
        # Save document
        filename = f'route_plan_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        doc.save(f'documents/routes/{filename}')
        return filename

    async def generate_customer_notification(self, 
                                          notification_type: str, 
                                          delivery_data: Dict[str, Any],
                                          customer_data: Dict[str, Any]) -> str:
        """Generate personalized customer notifications"""
        templates = {
            'delivery_confirmation': '''
                Dear {customer_name},
                
                Your delivery is scheduled for {delivery_date} between {time_window}.
                Tracking number: {tracking_number}
                
                You can track your delivery in real-time at: {tracking_url}
                
                Best regards,
                LogiSync Team
            ''',
            'delay_notification': '''
                Dear {customer_name},
                
                We apologize for the inconvenience, but your delivery scheduled for {delivery_date}
                has been delayed due to {delay_reason}.
                
                New estimated delivery time: {new_delivery_time}
                
                We are working to complete your delivery as soon as possible.
                
                Best regards,
                LogiSync Team
            ''',
            'proof_of_delivery': '''
                Dear {customer_name},
                
                This confirms that your delivery was completed on {delivery_date} at {delivery_time}.
                
                Signed by: {recipient_name}
                POD Reference: {pod_reference}
                
                Thank you for choosing our services.
                
                Best regards,
                LogiSync Team
            '''
        }
        
        # Use Ollama to personalize the message
        template = templates[notification_type]
        formatted_template = template.format(**delivery_data, **customer_data)
        
        prompt = f"""
        As a logistics communication expert, please personalize this notification while maintaining a professional tone:
        {formatted_template}
        
        Consider the customer's communication preferences and past interaction history.
        Make the message more engaging and personal while keeping it professional.
        """
        
        return await ollama.generate_text(prompt)

    async def generate_compliance_report(self, 
                                      report_type: str, 
                                      data: Dict[str, Any]) -> str:
        """Generate regulatory compliance reports"""
        if report_type == "FMCSA":
            return await self._generate_fmcsa_report(data)
        elif report_type == "safety_inspection":
            return await self._generate_safety_report(data)
        elif report_type == "environmental":
            return await self._generate_environmental_report(data)
        elif report_type == "driver_qualification":
            return await self._generate_driver_qualification_report(data)
        else:
            raise ValueError(f"Unsupported report type: {report_type}")

    async def _generate_fmcsa_report(self, data: Dict[str, Any]) -> str:
        """Generate FMCSA compliance report"""
        doc = SimpleDocTemplate(
            f'documents/compliance/fmcsa_{datetime.now().strftime("%Y%m%d")}.pdf',
            pagesize=letter
        )
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        story.append(Paragraph('FMCSA Compliance Report', styles['Title']))
        story.append(Spacer(1, 12))
        
        # Add carrier information
        story.append(Paragraph('Carrier Information', styles['Heading1']))
        story.append(Paragraph(f'Carrier: {data["carrier_name"]}', styles['Normal']))
        story.append(Paragraph(f'USDOT Number: {data["usdot_number"]}', styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Add operational data
        story.append(Paragraph('Operational Data', styles['Heading1']))
        for record in data["operational_records"]:
            story.append(Paragraph(
                f'Date: {record["date"]}\n'
                f'Miles Driven: {record["miles_driven"]}\n'
                f'Hours of Service: {record["hours_of_service"]}\n'
                f'Violations: {record["violations"]}\n',
                styles['Normal']
            ))
        
        doc.build(story)
        return doc.filename

    async def _generate_safety_report(self, data: Dict[str, Any]) -> str:
        """Generate safety inspection report"""
        doc = Document()
        doc.add_heading('Safety Inspection Report', 0)
        
        # Add inspection details
        doc.add_heading('Inspection Details', level=1)
        doc.add_paragraph(f'Inspector: {data["inspector_name"]}')
        doc.add_paragraph(f'Date: {data["inspection_date"]}')
        doc.add_paragraph(f'Location: {data["location"]}')
        
        # Add inspection results
        doc.add_heading('Inspection Results', level=1)
        for category, checks in data["inspection_items"].items():
            doc.add_heading(category, level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Item'
            header_cells[1].text = 'Status'
            header_cells[2].text = 'Notes'
            
            for check in checks:
                row_cells = table.add_row().cells
                row_cells[0].text = check["item"]
                row_cells[1].text = check["status"]
                row_cells[2].text = check["notes"]
        
        filename = f'safety_inspection_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        doc.save(f'documents/compliance/{filename}')
        return filename

    async def _generate_environmental_report(self, data: Dict[str, Any]) -> str:
        """Generate environmental impact report"""
        doc = SimpleDocTemplate(
            f'documents/compliance/environmental_{datetime.now().strftime("%Y%m%d")}.pdf',
            pagesize=letter
        )
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        story.append(Paragraph('Environmental Impact Assessment', styles['Title']))
        story.append(Spacer(1, 12))
        
        # Add summary
        story.append(Paragraph('Executive Summary', styles['Heading1']))
        story.append(Paragraph(data["executive_summary"], styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Add emissions data
        story.append(Paragraph('Emissions Data', styles['Heading1']))
        for emission_type, value in data["emissions_data"].items():
            story.append(Paragraph(
                f'{emission_type}: {value} metric tons CO2e',
                styles['Normal']
            ))
        
        # Add recommendations
        story.append(Paragraph('Recommendations', styles['Heading1']))
        for rec in data["recommendations"]:
            story.append(Paragraph(f'- {rec}', styles['Normal']))
        
        doc.build(story)
        return doc.filename

    async def _generate_driver_qualification_report(self, data: Dict[str, Any]) -> str:
        """Generate driver qualification report"""
        doc = Document()
        doc.add_heading('Driver Qualification File', 0)
        
        # Add driver information
        doc.add_heading('Driver Information', level=1)
        doc.add_paragraph(f'Name: {data["driver_name"]}')
        doc.add_paragraph(f'License Number: {data["license_number"]}')
        doc.add_paragraph(f'License Class: {data["license_class"]}')
        
        # Add qualification checklist
        doc.add_heading('Qualification Checklist', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Requirement'
        header_cells[1].text = 'Status'
        header_cells[2].text = 'Date'
        header_cells[3].text = 'Expiry'
        
        for item in data["qualification_items"]:
            row_cells = table.add_row().cells
            row_cells[0].text = item["requirement"]
            row_cells[1].text = item["status"]
            row_cells[2].text = item["date"]
            row_cells[3].text = item["expiry"]
        
        # Add violation history
        doc.add_heading('Violation History', level=1)
        if data["violations"]:
            for violation in data["violations"]:
                doc.add_paragraph(
                    f'Date: {violation["date"]}\n'
                    f'Type: {violation["type"]}\n'
                    f'Description: {violation["description"]}\n'
                    f'Resolution: {violation["resolution"]}'
                )
        else:
            doc.add_paragraph('No violations recorded.')
        
        filename = f'driver_qualification_{data["driver_name"]}_{datetime.now().strftime("%Y%m%d")}.docx'
        doc.save(f'documents/compliance/drivers/{filename}')
        return filename